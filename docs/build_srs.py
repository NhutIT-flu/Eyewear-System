from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/SRS_Eyewear_System.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
            set_cell_margins(row.cells[i])
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table, widths=None):
    set_table_borders(table)
    if widths:
        set_table_width(table, widths)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    format_table(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 20, "000000"),
        ("Heading 2", 16, "000000"),
        ("Heading 3", 14, "434343"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Software Requirements Specification")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Eyewear System UTH")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.bold = True

    meta = [
        ("Document type", "Software Requirements Specification (SRS)"),
        ("Project", "Eyewear System UTH"),
        ("Architecture", "Pure PHP API backend, MySQL database, HTML/CSS/JavaScript frontend"),
        ("Prepared for", "UTH software engineering coursework"),
        ("Version", "1.0"),
        ("Date", "26 May 2026"),
    ]
    doc.add_paragraph()
    add_table(doc, ["Field", "Value"], meta, [2.0, 4.5])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This document defines the functional and non-functional requirements for the Eyewear System based on the current repository structure, backend routes, database schema, and Jira module plan.")
    r.italic = True

    doc.add_page_break()


def add_revision_history(doc):
    doc.add_heading("Revision History", level=1)
    add_table(
        doc,
        ["Version", "Date", "Author", "Description"],
        [["1.0", "26 May 2026", "Project Team", "Initial SRS draft for Eyewear System"]],
        [1.0, 1.2, 1.5, 2.8],
    )


def add_introduction(doc):
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "The purpose of this SRS is to describe the requirements for the Eyewear System, an e-commerce and operations platform for selling eyewear products, managing prescriptions, processing orders, supporting staff workflows, and providing administrative analytics."
    )
    doc.add_heading("1.2 Scope", level=2)
    add_bullets(
        doc,
        [
            "Customer identity, registration, login, profile, addresses, and password recovery.",
            "Product catalog, categories, variants, lenses, images, inventory, and product search/filtering.",
            "Cart, voucher, checkout, order creation, order history, and payment processing.",
            "Sales support, prescription verification, customer support tickets, and return/refund workflows.",
            "Operations workflow for production, quality control, shipping, and dashboard reporting.",
            "Admin and manager functions for user management, reports, product statistics, and system configuration.",
        ],
    )
    doc.add_heading("1.3 Definitions and Abbreviations", level=2)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ("SRS", "Software Requirements Specification"),
            ("API", "Application Programming Interface"),
            ("RBAC", "Role-Based Access Control"),
            ("COD", "Cash on Delivery"),
            ("VNPay", "Vietnam online payment gateway"),
            ("SKU", "Stock Keeping Unit"),
            ("Prescription", "Optical prescription information such as sphere, cylinder, axis, and PD"),
        ],
        [1.5, 5.0],
    )


def add_overall_description(doc):
    doc.add_heading("2. Overall Description", level=1)
    doc.add_heading("2.1 Product Perspective", level=2)
    doc.add_paragraph(
        "The system is separated into a backend API and a frontend web interface. The backend follows an N-layered architecture with controllers, services, domain logic, infrastructure utilities, models, and a MySQL database. The frontend uses HTML, CSS, and JavaScript with modular pages, reusable components, layout loading, and API service modules."
    )
    doc.add_heading("2.2 User Classes", level=2)
    add_table(
        doc,
        ["Actor", "Description", "Main Responsibilities"],
        [
            ("Customer", "Public shopper using the web store", "Browse products, use virtual try-on, manage cart, checkout, view orders, submit support tickets"),
            ("Sales/Support Staff", "Staff responsible for customer and sales handling", "Verify orders, validate prescriptions, handle support tickets and returns"),
            ("Operations Staff", "Staff responsible for fulfillment", "Advance production steps, create shipments, update tracking"),
            ("Manager", "Business supervisor", "View reports, monitor revenue, product statistics, inventory, and team performance"),
            ("System Admin", "System owner/administrator", "Manage users, staff accounts, roles, permissions, and configuration"),
        ],
        [1.25, 2.2, 3.05],
    )
    doc.add_heading("2.3 Operating Environment", level=2)
    add_bullets(
        doc,
        [
            "Backend runtime: PHP 8.x using a custom MVC/router structure.",
            "Database: MySQL or MariaDB, commonly through XAMPP or Laragon.",
            "Frontend runtime: Browser-based HTML/CSS/JavaScript served by a local web server.",
            "API format: JSON over HTTP.",
            "Development tools: Git, Postman, and optionally Jenkins for automated build/test pipelines.",
        ],
    )
    doc.add_heading("2.4 Constraints", level=2)
    add_bullets(
        doc,
        [
            "The implementation must follow the existing Controller-Service-Model project structure.",
            "Protected API routes must use token authentication and permission checks.",
            "The database schema must remain compatible with the current MySQL tables.",
            "Payment integration may use mock/internal payment flow unless VNPay production credentials are available.",
            "Email features require valid SMTP configuration.",
        ],
    )


def add_functional_requirements(doc):
    doc.add_heading("3. Functional Requirements", level=1)
    modules = [
        (
            "3.1 Auth and Identity Module",
            [
                ("FR-AUTH-01", "The system shall allow users to register with name, email, and password."),
                ("FR-AUTH-02", "The system shall create customer accounts with a default customer role."),
                ("FR-AUTH-03", "The system shall generate email verification tokens for new accounts."),
                ("FR-AUTH-04", "The system shall allow active users to log in with valid credentials."),
                ("FR-AUTH-05", "The system shall return user roles, permissions, and an access token after login."),
                ("FR-AUTH-06", "The system shall prevent inactive users from logging in before verification."),
                ("FR-AUTH-07", "The system shall allow authenticated users to view and update their profile."),
                ("FR-AUTH-08", "The system shall allow authenticated users to change password after confirming current password."),
                ("FR-AUTH-09", "The system shall support forgot password and reset password flow using reset tokens."),
                ("FR-AUTH-10", "The system shall restrict protected APIs using role and permission authorization."),
            ],
        ),
        (
            "3.2 Product Catalog Module",
            [
                ("FR-PROD-01", "The system shall display a paginated list of active products."),
                ("FR-PROD-02", "The system shall display product detail by numeric ID or slug."),
                ("FR-PROD-03", "The system shall include variants, stock, image URL, price range, and category details in product responses."),
                ("FR-PROD-04", "The system shall allow product search by name, model, slug, or brand."),
                ("FR-PROD-05", "The system shall filter products by category, brand, gender, price range, and stock state."),
                ("FR-PROD-06", "The system shall provide product category and brand lists."),
                ("FR-PROD-07", "The system shall allow authorized staff to create, update, and deactivate products."),
                ("FR-PROD-08", "The system shall allow authorized staff to view and update inventory stock."),
                ("FR-PROD-09", "The system shall support product image URL management or image upload for product variants."),
            ],
        ),
        (
            "3.3 Cart and Checkout Module",
            [
                ("FR-CART-01", "The system shall allow authenticated customers to add product variants to cart."),
                ("FR-CART-02", "The system shall validate stock before adding or updating cart quantity."),
                ("FR-CART-03", "The system shall allow customers to update item quantity and remove cart items."),
                ("FR-CART-04", "The system shall allow customers to select or deselect cart items for checkout."),
                ("FR-CART-05", "The system shall calculate subtotal, discount, shipping fee, and final total."),
                ("FR-CART-06", "The system shall support valid voucher application and removal."),
                ("FR-CART-07", "The system shall reject invalid, inactive, or expired vouchers."),
                ("FR-CART-08", "The system shall validate selected items, shipping address, preorder rules, and stock before checkout."),
                ("FR-CART-09", "The system shall reduce stock and clear purchased cart items after successful checkout."),
            ],
        ),
        (
            "3.4 Order and Payment Module",
            [
                ("FR-PAY-01", "The system shall create orders from selected cart items."),
                ("FR-PAY-02", "The system shall create order items and initialize a payment record for each order."),
                ("FR-PAY-03", "The system shall allow customers to view their order history and order detail."),
                ("FR-PAY-04", "The system shall process COD, bank transfer, card, or e-wallet payment methods."),
                ("FR-PAY-05", "The system shall store payment status, transaction reference, and paid timestamp."),
                ("FR-PAY-06", "The system shall allow authorized staff to confirm pending payments."),
                ("FR-PAY-07", "The system shall support VNPay payment URL generation and return handling when payment gateway credentials are configured."),
                ("FR-PAY-08", "The system shall support refund workflow for eligible paid orders."),
            ],
        ),
        (
            "3.5 Sales, Support, Operations, and Admin Module",
            [
                ("FR-OPS-01", "The system shall allow sales staff to view and verify customer orders."),
                ("FR-OPS-02", "The system shall allow staff to manage support tickets and replies."),
                ("FR-OPS-03", "The system shall allow operations staff to advance production workflow steps."),
                ("FR-OPS-04", "The system shall allow operations staff to create and update shipment records."),
                ("FR-OPS-05", "The system shall provide an admin dashboard with revenue, active orders, paid orders, top products, and top categories."),
                ("FR-OPS-06", "The system shall allow admins to manage users and staff accounts."),
                ("FR-OPS-07", "The system shall provide revenue analytics and sales reports by day range."),
                ("FR-OPS-08", "The system shall provide product statistics for top-selling products and categories."),
                ("FR-OPS-09", "The system shall support notification creation, listing, and read status."),
                ("FR-OPS-10", "The system shall support report export in CSV or Excel-compatible format."),
            ],
        ),
    ]

    for heading, rows in modules:
        doc.add_heading(heading, level=2)
        add_table(doc, ["ID", "Requirement"], rows, [1.25, 5.25])


def add_nonfunctional_requirements(doc):
    doc.add_heading("4. Non-Functional Requirements", level=1)
    rows = [
        ("NFR-01", "Security", "Passwords shall be hashed before storage. Protected APIs shall require bearer token authentication."),
        ("NFR-02", "Authorization", "Role and permission checks shall be enforced for staff and admin features."),
        ("NFR-03", "Data Integrity", "Checkout and order creation shall use database transactions where stock, order, payment, and cart changes must stay consistent."),
        ("NFR-04", "Usability", "Customer and staff pages shall be clear, responsive, and easy to navigate."),
        ("NFR-05", "Maintainability", "Code shall follow the current Controller-Service-Model and modular frontend service structure."),
        ("NFR-06", "Performance", "Product listing and dashboard endpoints shall return paginated or summarized data to avoid excessive response size."),
        ("NFR-07", "Compatibility", "The system shall run on PHP 8.x, MySQL/MariaDB, and modern browsers."),
        ("NFR-08", "Reliability", "API errors shall return consistent JSON responses with clear status codes and messages."),
        ("NFR-09", "Auditability", "Payment records shall include transaction references and timestamps for tracking."),
        ("NFR-10", "Testability", "Major APIs shall be verifiable using Postman collection or backend feature tests."),
    ]
    add_table(doc, ["ID", "Category", "Requirement"], rows, [1.0, 1.4, 4.1])


def add_external_interfaces(doc):
    doc.add_heading("5. External Interface Requirements", level=1)
    doc.add_heading("5.1 User Interface", level=2)
    add_bullets(
        doc,
        [
            "Customer pages: home, shop, details, cart, checkout, auth, profile, wishlist, support.",
            "Staff/admin portal pages: overview, products, inventory, orders, users, analytics, operations, promotions, support, settings.",
            "Shared layouts: header, footer, customer sidebar, staff sidebar, and modal components.",
        ],
    )
    doc.add_heading("5.2 API Interface", level=2)
    add_table(
        doc,
        ["Module", "Representative Endpoints"],
        [
            ("Auth", "POST /api/v1/auth/register, POST /api/v1/auth/login, POST /api/v1/auth/logout, GET /api/v1/auth/me"),
            ("Product", "GET /api/v1/products, GET /api/v1/products/{id}, GET /api/v1/products/categories"),
            ("Cart", "GET /api/v1/cart, POST /api/v1/cart, PUT /api/v1/cart/items/{id}, DELETE /api/v1/cart/items/{id}"),
            ("Checkout", "POST /api/v1/checkout"),
            ("Payment", "POST /api/v1/payments/process, GET /api/v1/payments/status, POST /api/v1/payments/confirm"),
            ("Admin", "GET /api/v1/admin/users, POST /api/v1/admin/staff, GET /api/v1/dashboard"),
            ("Operations", "GET /api/v1/ops, POST /api/v1/ops/advance, POST /api/v1/ops/shipments"),
        ],
        [1.4, 5.1],
    )
    doc.add_heading("5.3 Hardware and Software Interfaces", level=2)
    add_bullets(
        doc,
        [
            "MySQL database server for persistence.",
            "SMTP mail server for email verification and password reset.",
            "VNPay gateway for online payment when enabled.",
            "Postman for manual API verification.",
            "Jenkins may be used for automated test and deployment workflows.",
        ],
    )


def add_data_requirements(doc):
    doc.add_heading("6. Data Requirements", level=1)
    add_table(
        doc,
        ["Domain", "Main Entities"],
        [
            ("Identity", "role, user, user_roles, permissions, role_permissions, profiles, user_addresses"),
            ("Catalog", "category, product, productvariant, inventory, lens, promotion"),
            ("Commerce", "prescription, cart, cartitem, wishlist"),
            ("Orders", "order, orderitem, payment, shipment"),
            ("Support", "supportticket, ticket_replies, returnrequest"),
        ],
        [1.4, 5.1],
    )
    doc.add_paragraph(
        "The database shall enforce relationships between users, roles, products, variants, carts, orders, payments, and support records using primary keys, foreign keys, and unique constraints where appropriate."
    )


def add_use_cases(doc):
    doc.add_heading("7. Use Cases", level=1)
    use_cases = [
        ("UC-01", "Register Account", "Customer", "Customer submits name, email, and password; system creates inactive account and sends verification link."),
        ("UC-02", "Login", "Customer/Staff/Admin", "User submits credentials; system validates account and returns user data, roles, permissions, and token."),
        ("UC-03", "Browse Product Catalog", "Customer", "Customer views paginated products and filters by category, brand, gender, price, or stock."),
        ("UC-04", "Manage Cart", "Customer", "Customer adds products, updates quantity, applies voucher, and reviews totals."),
        ("UC-05", "Checkout Order", "Customer", "Customer checks out selected cart items; system validates stock, creates order, payment, and order items."),
        ("UC-06", "Confirm Payment", "Sales/Manager", "Authorized staff confirms pending payment and system updates payment/order status."),
        ("UC-07", "Manage Products", "Admin/Manager", "Authorized staff creates, updates, deactivates products and updates inventory."),
        ("UC-08", "Process Shipment", "Operations Staff", "Operations staff advances production and creates/updates shipment tracking."),
        ("UC-09", "View Dashboard", "Manager/Admin", "Manager views revenue, active orders, top products, and sales reports."),
        ("UC-10", "Handle Support Ticket", "Customer/Support Staff", "Customer submits ticket; staff replies and updates ticket status."),
    ]
    add_table(doc, ["ID", "Use Case", "Primary Actor", "Summary"], use_cases, [0.75, 1.35, 1.4, 3.0])


def add_test_traceability(doc):
    doc.add_heading("8. Test Case Traceability", level=1)
    rows = [
        ("TC-AUTH-01", "Login Success", "FR-AUTH-04, FR-AUTH-05"),
        ("TC-AUTH-02", "Wrong Password", "FR-AUTH-04"),
        ("TC-AUTH-03", "Missing Email or Password", "FR-AUTH-04"),
        ("TC-AUTH-04", "Register Success", "FR-AUTH-01, FR-AUTH-02, FR-AUTH-03"),
        ("TC-AUTH-05", "Duplicate Email", "FR-AUTH-01"),
        ("TC-AUTH-06", "Unauthorized Access", "FR-AUTH-10"),
        ("TC-PROD-01", "Product Listing", "FR-PROD-01"),
        ("TC-PROD-02", "Product Detail", "FR-PROD-02, FR-PROD-03"),
        ("TC-PROD-03", "Product Search and Filtering", "FR-PROD-04, FR-PROD-05"),
        ("TC-PROD-04", "Admin Product CRUD", "FR-PROD-07"),
        ("TC-CART-01", "Add To Cart", "FR-CART-01, FR-CART-02"),
        ("TC-CART-02", "Remove From Cart", "FR-CART-03"),
        ("TC-CART-03", "Update Quantity", "FR-CART-02, FR-CART-03"),
        ("TC-CART-04", "Voucher Discount", "FR-CART-06, FR-CART-07"),
        ("TC-CART-05", "Checkout Validation", "FR-CART-08, FR-CART-09"),
        ("TC-PAY-01", "Order Creation", "FR-PAY-01, FR-PAY-02"),
        ("TC-PAY-02", "Payment Validation", "FR-PAY-04, FR-PAY-05, FR-PAY-06"),
        ("TC-PAY-03", "VNPay Success and Failure", "FR-PAY-07"),
        ("TC-PAY-04", "Refund API", "FR-PAY-08"),
        ("TC-OPS-01", "Admin Dashboard", "FR-OPS-05"),
        ("TC-OPS-02", "User Management", "FR-OPS-06"),
        ("TC-OPS-03", "Revenue Analytics", "FR-OPS-07"),
        ("TC-OPS-04", "Product Statistics", "FR-OPS-08"),
        ("TC-OPS-05", "Notification System", "FR-OPS-09"),
        ("TC-OPS-06", "Report Export", "FR-OPS-10"),
    ]
    add_table(doc, ["Test ID", "Test Case", "Related Requirement"], rows, [1.1, 2.4, 3.0])

    doc.add_heading("8.1 Definition of Done", level=2)
    add_bullets(
        doc,
        [
            "The related API or feature is implemented in the existing architecture.",
            "Valid and invalid cases are tested through Postman or backend test scripts.",
            "Expected and actual status codes are recorded.",
            "Response evidence or screenshots are attached to the Jira ticket.",
            "The ticket is moved to Done only after the implemented behavior passes verification.",
        ],
    )


def add_appendix(doc):
    doc.add_heading("9. Appendix: Jira Module Mapping", level=1)
    add_table(
        doc,
        ["Epic", "Stories"],
        [
            ("ESQ-1: Auth Module", "Login, Register, User Profile"),
            ("ESQ-2: Product Module", "Product Catalog, Product Details"),
            ("ESQ-3: Cart & Checkout", "Manage Shopping Cart, Checkout Process"),
            ("ESQ-4: Order & Payment", "Process Order Payment, Confirm Payment (Staff), View Order History"),
            ("ESQ-5: Operations", "Inventory Management"),
            ("ESQ-6: Dashboard & Admin", "Sales Dashboard, User Management"),
        ],
        [1.7, 4.8],
    )


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    set_styles(doc)
    add_title_page(doc)
    add_revision_history(doc)
    add_introduction(doc)
    add_overall_description(doc)
    add_functional_requirements(doc)
    add_nonfunctional_requirements(doc)
    add_external_interfaces(doc)
    add_data_requirements(doc)
    add_use_cases(doc)
    add_test_traceability(doc)
    add_appendix(doc)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
